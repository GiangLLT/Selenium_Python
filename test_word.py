from docx import Document
import os
from datetime import datetime

def replace_text_with_format(paragraph, key, value):
    # Tạo một danh sách mới để chứa các đoạn văn bản đã thay thế
    new_paragraph_runs = []

    for run in paragraph.runs:
        if key in run.text:
            # Chia nhỏ các đoạn văn bản bằng khóa
            parts = run.text.split(key)
            for i, part in enumerate(parts):
                # Tạo một đoạn văn bản mới với phần trước và sau khóa
                if part:
                    new_run = run._element
                    new_run.text = part
                    new_paragraph_runs.append(new_run)
                if i < len(parts) - 1:
                    # Chèn giá trị thay thế với định dạng tương tự
                    replacement_run = run._element
                    replacement_run.text = value
                    new_paragraph_runs.append(replacement_run)
        else:
            # Giữ nguyên đoạn văn bản nếu không chứa khóa
            new_paragraph_runs.append(run._element)

    # Xóa tất cả các đoạn văn bản hiện có trong đoạn văn
    while paragraph.runs:
        paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)

    # Thêm các đoạn văn bản mới vào đoạn văn
    for run_element in new_paragraph_runs:
        paragraph._element.append(run_element)

def write_to_word(template_path, output_directory, data):
    # Mở file Word template
    doc = Document(template_path)

    # Duyệt qua tất cả các đoạn văn trong tài liệu và thay thế placeholders
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                if any(run.bold or run.italic or run.underline or run.font.name or (run.font.size is not None) for run in paragraph.runs):
                    replace_text_with_format(paragraph, f'{{{key}}}', value)
                else:
                    paragraph.text = paragraph.text.replace(f'{{{key}}}', value)

    # Duyệt qua tất cả các bảng trong tài liệu và thay thế placeholders
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        if key in paragraph.text:
                            if any(run.bold or run.italic or run.underline or run.font.name or (run.font.size is not None) for run in paragraph.runs):
                                replace_text_with_format(paragraph, f'{{{key}}}', value)
                            else:
                                paragraph.text = paragraph.text.replace(f'{{{key}}}', value)
    
    # Định dạng lại tên file với ngày tháng năm giờ phút
    current_time = datetime.now().strftime("%d%m%Y_%H%M")
    output_filename = f"Survey_{current_time}.docx"
    output_path = os.path.join(output_directory, output_filename)
    
    # Lưu file Word mới với dữ liệu đã thay thế
    doc.save(output_path)
    
    print(f"File đã được ghi dữ liệu và lưu tại: {output_path}")
    return output_path

if __name__ == "__main__":
    # Đọc dữ liệu từ tệp Excel
    data = {
        "CM_RATE0": "0/23",
        "CM_VALUE0": "0.0%",
        "CM_RATE1": "9/23",
        "CM_VALUE1": "39.13%",
        "CM_RATE2": "9/23",
        "CM_VALUE2": "39.13%",
        "CM_RATE3": "5/23",
        "CM_VALUE3": "21.74%",
        "CM_VALUE4": "78.26%",
        "CM_VALUE5": "100.0%",
        # "LD_RATE0": "1/23",
        # "LD_VALUE0": "4.35%",
        # "LD_RATE1": "13/23",
        # "LD_VALUE1": "56.52%",
        # "LD_RATE2": "4/23",
        # "LD_VALUE2": "17.39%",
        # "LD_RATE3": "5/23",
        # "LD_VALUE3": "21.74%",
        # "LD_VALUE4": "73.91%",
        # "LD_VALUE5": "94.44%",
        # "DD_RATE0": "0/23",
        # "DD_VALUE0": "0.0%",
        # "DD_RATE1": "9/23",
        # "DD_VALUE1": "39.13%",
        # "DD_RATE2": "9/23",
        # "DD_VALUE2": "39.13%",
        # "DD_RATE3": "5/23",
        # "DD_VALUE3": "21.74%",
        # "DD_VALUE4": "78.26%",
        # "DD_VALUE5": "100.0%",
        # "DU_RATE0": "0/23",
        # "DU_VALUE0": "0.0%",
        # "DU_RATE1": "11/23",
        # "DU_VALUE1": "47.83%",
        # "DU_RATE2": "7/23",
        # "DU_VALUE2": "30.43%",
        # "DU_RATE3": "5/23",
        # "DU_VALUE3": "21.74%",
        # "DU_VALUE4": "78.26%",
        # "DU_VALUE5": "100.0%",
        # "KL_RATE0": "0/23",
        # "KL_VALUE0": "0.0%",
        # "KL_RATE1": "9/23",
        # "KL_VALUE1": "39.13%",
        # "KL_RATE2": "9/23",
        # "KL_VALUE2": "39.13%",
        # "KL_RATE3": "5/23",
        # "KL_VALUE3": "21.74%",
        # "KL_VALUE4": "78.26%",
        # "KL_VALUE5": "100.0%",
        # "NH_RATE0": "0/23",
        # "NH_VALUE0": "0.0%",
        # "NH_RATE1": "8/23",
        # "NH_VALUE1": "34.78%",
        # "NH_RATE2": "10/23",
        # "NH_VALUE2": "43.48%",
        # "NH_RATE3": "5/23",
        # "NH_VALUE3": "21.74%",
        # "NH_VALUE4": "78.26%",
        # "NH_VALUE5": "100.0%",
        # "LN_RATE0": "1/23",
        # "LN_VALUE0": "4.35%",
        # "LN_RATE1": "12/23",
        # "LN_VALUE1": "52.17%",
        # "LN_RATE2": "5/23",
        # "LN_VALUE2": "21.74%",
        # "LN_RATE3": "5/23",
        # "LN_VALUE3": "21.74%",
        # "LN_VALUE4": "73.91%",
        # "LN_VALUE5": "94.44%"
    }

    try:
        # Đường dẫn tệp nguồn
        src_file_template = ".\\Template\\Template_Survey.docx"
        dest_directory = ".\\Survey"
        src_output_survey = write_to_word(src_file_template, dest_directory, data)
    except ValueError:
        print("Error")

    input("Nhấn phím bất kỳ để kết thúc...")
