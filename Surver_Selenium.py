from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import ElementClickInterceptedException
import pandas as pd
import time
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.firefox.options import Options
from selenium.webdriver.edge.options import Options
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
import os
import shutil
from datetime import datetime
from docx import Document
from docx2pdf import convert
import fitz  # PyMuPDF
import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk
import os
import json

def RPA_workit(data_excel):
    list_data = { 
        'No Data': 0,
        'Tắt kích hoạt Thành Công': 0,
        'Đã Tắt kích hoạt': 0,
        'Kích hoạt thành công': 0,
        'Đã  Kích hoạt': 0,
    }
    data_Nodata = list_data["No Data"]
    data_Unactive = list_data["Tắt kích hoạt Thành Công"]
    data_Unactived = list_data["Đã Tắt kích hoạt"]
    data_Active = list_data["Kích hoạt thành công"]
    data_Actived = list_data["Đã  Kích hoạt"]

    # Chuyển DataFrame thành list các dòng (records)
    records = data_excel.to_dict(orient='records')

    driver = setup_firefox_driver()
    # driver = webdriver.Firefox()

    try:
        # Mở trang web
        driver.get("https://eoffice.bamboocap.com.vn/")

        # Tìm các trường văn bản bằng ID và nhập dữ liệu
        username = driver.find_element(By.ID, "taikhoan")
        password = driver.find_element(By.ID, "matkhau")
        buttonLogin = driver.find_element(By.ID, "btn_Login")

        username.send_keys("INTEGRATION@BAMBOO.COM.VN")
        password.send_keys("Abc@123")
        buttonLogin.click()

        # Đợi cho đến khi trang chuyển hướng sau khi đăng nhập
        WebDriverWait(driver, 10).until(EC.url_contains("LoadContent"))

        print("Đăng nhập web thành công!!!")

        # Mở trang sau khi đăng nhập
        driver.get("https://eoffice.bamboocap.com.vn/LoadContent?rightcode=9502&v=23")
        WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, '//th[contains(@data-title, "Mã đăng nhập")]/a')))

        print("Bắt đầu xử lý.....")
        for data in records:
            # list_values = list(list_data.values())
            # print(list_values)

            data_stt = str(data["STT"])
            data_email = data["Email"]
            data_status = data["Status"]
            
            time.sleep(1)
            # Tìm và nhấp vào thẻ <a> bên trong <th> có thuộc tính data-title cụ thể
            link = driver.find_element(By.XPATH, '//th[contains(@data-title, "Mã đăng nhập")]/a')
            # Cuộn đến phần tử để đảm bảo nó nằm trong khung nhìn
            driver.execute_script("arguments[0].scrollIntoView();", link)
            WebDriverWait(driver, 10).until(EC.element_to_be_clickable(link))

            # Thử nhấp vào phần tử, nếu bị chặn, đợi và thử lại
            try:
                link.click()
            except ElementClickInterceptedException:
                print("Phần tử bị chặn, thử lại...")
                try:
                    button_save_rep = driver.find_element(By.XPATH, '//a[@class="k-button k-button-icontext k-primary k-grid-update"]')
                    button_save_rep.click()
                    time.sleep(1)
                    link.click()
                except Exception as e:
                    time.sleep(1)
                    link.click()

            # Đợi phần tử textbox và button xuất hiện
            WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.CLASS_NAME, "k-textbox")))
            filter_value = driver.find_element(By.CLASS_NAME, "k-textbox")
            filter_button = driver.find_element(By.CLASS_NAME, "k-primary")

            # Cuộn đến phần tử để đảm bảo nó nằm trong khung nhìn
            driver.execute_script("arguments[0].scrollIntoView();", filter_value)

            # Đợi cho đến khi phần tử có thể tương tác được
            WebDriverWait(driver, 10).until(EC.element_to_be_clickable((By.CLASS_NAME, "k-textbox")))

            filter_value.clear()
            filter_value.send_keys(data_email)
            filter_button.click()

            try:
                WebDriverWait(driver, 3).until(
                    EC.presence_of_element_located((By.XPATH, f'//td[text()="{data_email.upper()}"]'))
                )

                # Đợi và nhấp vào nút chỉnh sửa
                WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.NAME, 'edit')))
                edit_button = driver.find_element(By.NAME, "edit")
                edit_button.click()

                # Đợi phần tử checkbox và các nút lưu/hủy xuất hiện
                WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.ID, "check_active")))
                checkbox_status = driver.find_element(By.ID, "check_active")
                button_save = driver.find_element(By.XPATH, '//a[@class="k-button k-button-icontext k-primary k-grid-update"]')
                button_cancel = driver.find_element(By.XPATH, '//a[@class="k-button k-button-icontext k-grid-cancel"]')

                if data_status == False:
                    if checkbox_status.is_selected():
                        checkbox_status.click()
                        button_save.click()
                        data_Unactive += 1
                        # list_data["Tắt kích hoạt Thành Công"] = data_Unactive
                        print(f"======================================================")
                        print(f"{data_stt} || {data_email} || Tắt kích hoạt Thành Công")
                    else:
                        button_cancel.click()
                        data_Unactived += 1
                        # list_data["Đã Tắt kích hoạt"] = data_Unactived
                        print(f"======================================================")
                        print(f"{data_stt} || {data_email} || Đã Tắt kích hoạt")
                else:
                    if not checkbox_status.is_selected():
                        checkbox_status.click()
                        button_save.click()
                        data_Active += 1
                        # list_data["Kích hoạt thành công"] = data_Active
                        print(f"======================================================")
                        print(f"{data_stt} || {data_email} || Kích hoạt thành công")
                    else:
                        button_cancel.click()
                        data_Actived += 1
                        # list_data["Đã  Kích hoạt"] = data_Actived
                        print(f"======================================================")
                        print(f"{data_stt} || {data_email} || Đã Kích hoạt")
            except Exception as e:
                data_Nodata += 1
                # list_data["No Data"] = data_Nodata
                print(f"======================================================")
                print(f"{data_stt} || {data_email} || No data")
                continue
        
        list_data["Tắt kích hoạt Thành Công"] = data_Unactive
        list_data["Đã Tắt kích hoạt"] = data_Unactived
        list_data["Kích hoạt thành công"] = data_Active
        list_data["Đã  Kích hoạt"] = data_Actived
        list_data["No Data"] = data_Nodata
        print(list_data)
    finally:
        # Đóng trình duyệt
        driver.quit()

def read_excel(file_path, sheet_name=0):
    """
    Đọc dữ liệu từ tệp Excel và trả về DataFrame.

    :param file_path: Đường dẫn tới tệp Excel.
    :param sheet_name: Tên sheet hoặc chỉ số sheet (mặc định là 0 - sheet đầu tiên).
    :return: DataFrame chứa dữ liệu từ tệp Excel.
    """
    try:
        # Đọc dữ liệu từ tệp Excel
        df = pd.read_excel(file_path, sheet_name=sheet_name)
        return df
    except Exception as e:
        print(f"Đã xảy ra lỗi khi đọc tệp Excel: {e}")
        return None

def setup_chrome_driver():
    chrome_options = webdriver.ChromeOptions()
    chrome_options.add_argument('--headless')
    chrome_options.add_argument('--disable-gpu')
    chrome_options.add_argument('--no-sandbox')
    chrome_options.add_argument('--disable-dev-shm-usage')

    driver = webdriver.Chrome(options=chrome_options)
    return driver

def setup_firefox_driver():
    firefox_options = webdriver.FirefoxOptions()
    firefox_options.add_argument('--headless')
    firefox_options.add_argument('--disable-gpu')
    firefox_options.add_argument('--no-sandbox')
    firefox_options.add_argument('--disable-dev-shm-usage')
    firefox_options.add_argument('--private')

    driver = webdriver.Firefox(options=firefox_options)
    return driver

def setup_edge_driver():
    edge_options =webdriver.EdgeOptions()
    edge_options.add_argument('--headless')
    edge_options.add_argument('--disable-gpu')
    edge_options.add_argument('--no-sandbox')
    edge_options.add_argument('--disable-dev-shm-usage')
    edge_options.add_argument('--inprivate')

    driver = webdriver.Edge(options=edge_options)
    return driver

def login_to_website(driver, Link_Web):
    driver.get(Link_Web)
    
    print("Login Starting.....")
    # Chờ và chuyển vào iframe hrdIframe
    iframe = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, '//iframe[@id="hrdIframe"]')))
    driver.switch_to.frame(iframe)
    
    # Nhập username và click Next''
    time.sleep(5)
    username = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, '//input[@type="email"]')))
    username.send_keys("it@bamboocap.com.vn")
    btnNext = driver.find_element(By.XPATH, '//input[@value="Next"]')
    btnNext.click()
    
    time.sleep(3)
    # Nhập password và click Sign in
    password = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, '//input[@type="password"]')))
    password.send_keys("#BcG2022!")
    btnSignIn = driver.find_element(By.XPATH, '//input[@value="Sign in"]')
    btnSignIn.click()
    
    # Chờ và click Yes
    btnYes = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, '//input[@value="Yes"]')))
    btnYes.click()
    
    # Chuyển ra khỏi iframe về default content
    driver.switch_to.default_content()
    print("Login Success.....")

def read_survey_new(driver):
    list_survey = [
        {
            "Vui lòng đánh giá Năng lực chuyên môn?": {
                "Chưa Tốt": 0,
                "Tốt": 0,
                "Rất Tốt": 0,
            }
        },
        {
            "Vui lòng đánh giá Năng lực lãnh đạo?": {
                "Chưa Tốt": 0,
                "Tốt": 0,
                "Rất Tốt": 0,
            }
        },
        {
            "Vui lòng đánh giá Năng lực làm việc đồng đội?": {
                "Chưa Tốt": 0,
                "Tốt": 0,
                "Rất Tốt": 0,
            }
        },
        {
            "Đạo đức (Liêm chính; Chính trực; Trung thành; Tận tuỵ; Lòng biết ơn)": {
                "Chưa Tốt": 0,
                "Tốt": 0,
                "Rất Tốt": 0,
            }
        },
        {
            "Tuân thủ kỷ luật (Tuân thủ Nội quy lao động; Tuân thủ Điều lệ; Quy chế quản lý nội bộ; Quy trình làm việc)": {
                "Chưa Tốt": 0,
                "Tốt": 0,
                "Rất Tốt": 0,
            }
        },
        {
            "Nhiệt huyết": {
                "Chưa Tốt": 0,
                "Tốt": 0,
                "Rất Tốt": 0,
            }
        },
        {
            "Lắng nghe, tôn trọng và đối xử công bằng": {
                "Chưa Tốt": 0,
                "Tốt": 0,
                "Rất Tốt": 0,
            }
        }
    ]
    # WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.CLASS_NAME, 'suiteheader')))
    # title_form = driver.find_element(By.XPATH, '//div[@role="heading"]').text
    # print("Survey Title:", title_form)        

def read_survey_final(driver,all_reponse = 16):
    print("Proceess Loading....")
    list_survey = [
        {
            "Vui lòng đánh giá Năng lực chuyên môn?": {
                "Chưa Tốt": 0,
                "Tốt": 0,
                "Rất Tốt": 0,
            }
        },
        {
            "Vui lòng đánh giá Năng lực lãnh đạo?": {
                "Chưa Tốt": 0,
                "Tốt": 0,
                "Rất Tốt": 0,
            }
        },
        {
            "Vui lòng đánh giá Năng lực làm việc đồng đội?": {
                "Chưa Tốt": 0,
                "Tốt": 0,
                "Rất Tốt": 0,
            }
        },
        {
            "Đạo đức (Liêm chính; Chính trực; Trung thành; Tận tuỵ; Lòng biết ơn)": {
                "Chưa Tốt": 0,
                "Tốt": 0,
                "Rất Tốt": 0,
            }
        },
        {
            "Tuân thủ kỷ luật (Tuân thủ Nội quy lao động; Tuân thủ Điều lệ; Quy chế quản lý nội bộ; Quy trình làm việc)": {
                "Chưa Tốt": 0,
                "Tốt": 0,
                "Rất Tốt": 0,
            }
        },
        {
            "Nhiệt huyết": {
                "Chưa Tốt": 0,
                "Tốt": 0,
                "Rất Tốt": 0,
            }
        },
        {
            "Lắng nghe, tôn trọng và đối xử công bằng": {
                "Chưa Tốt": 0,
                "Tốt": 0,
                "Rất Tốt": 0,
            }
        }
    ]

    # list_survey = [
    #     {
    #         "Vui lòng đánh giá Năng lực chuyên môn?": {
    #             "Chưa Tốt": 0,
    #             "Tốt": 0,
    #             "Rất Tốt": 0,
    #         }
    #     },
    #     {
    #         "Vui lòng đánh giá Năng lực lãnh đạo?": {
    #             "Chưa Tốt": 0,
    #             "Tốt": 0,
    #             "Rất Tốt": 0,
    #         }
    #     },
    #     {
    #         "Vui lòng đánh giá Năng lực làm việc đồng đội?": {
    #             "Chưa Tốt": 0,
    #             "Tốt": 0,
    #             "Rất Tốt": 0,
    #         }
    #     },
    #     {
    #         "Đạo đức (Liêm chính; Chính trực; Trung thành; Tận tuỵ; Lòng biết ơn)": {
    #             "Chưa Tốt": 0,
    #             "Tốt": 0,
    #             "Rất Tốt": 0,
    #         }
    #     },
    #     {
    #         "Tuân thủ kỷ luật (Tuân thủ Nội quy lao động; Tuân thủ Điều lệ; Quy chế quản lý nội bộ; Quy trình làm việc)": {
    #             "Chưa Tốt": 0,
    #             "Tốt": 0,
    #             "Rất Tốt": 0,
    #         }
    #     },
    #     {
    #         "Nhiệt huyết": {
    #             "Chưa Tốt": 0,
    #             "Tốt": 0,
    #             "Rất Tốt": 0,
    #         }
    #     },
    #     {
    #         "Lắng nghe, tôn trọng và đối xử công bằng": {
    #             "Chưa Tốt": 0,
    #             "Tốt": 0,
    #             "Rất Tốt": 0,
    #         }
    #     }
    # ]

    # Duyệt qua từng danh mục và in tên của từng danh mục
    time.sleep(3)
    response_element = WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.XPATH, '//div[contains(@aria-label, "responses")]'))
    )
    # Lấy giá trị số lượng responses
    # response_value = response_element.text.strip()
    response_value = response_element.text.strip().split('\n')[-1]
    empty_value = all_reponse - int(response_value)
    print(f'Số lượng tham gia khảo Sát:{response_value}')
    # label_t = "Tốt"
    # label_rt ="Rất Tốt"
    for survey_item in list_survey:
        good_value = 0
        good_rate = 0
        for key in survey_item:
            print(key)
            print("------------------------------------------------------------")
            xpath_query = f'//div[contains(@aria-label, "{key}")]//tbody//tr'
            rows = driver.find_elements(By.XPATH, xpath_query)
        
            for row in rows:
                label = row.find_element(By.CLASS_NAME, 'chart-control-legend-label').text
                value = row.find_element(By.CLASS_NAME, 'chart-control-legend-value').text
                vote_value = int(value)/all_reponse*100
                print(f"{label}: {value} - {vote_value}")
                # if label == "Tốt" or label == "Rất Tốt" or label == "tốt" or label == "Rất tốt":
                # if label.lower() == label_t.lower() or label.lower() == label_rt.lower():
                if label.lower() in {"tốt", "rất tốt"}:
                    good_value += vote_value
                    good_rate += int(value)
               
        lable_empty = empty_value/all_reponse*100
        print(f"Không Khảo Sát: {empty_value} - {lable_empty}%")

        print(f"Tỷ lệ Rất tốt và Tốt/Tổng số lượng mời tham gia khảo sát: {good_value}%")
        
        label_good = good_rate/int(response_value)*100
        label_good_rounded = round(label_good, 2)
        print(f"Tỷ lệ Rất tốt và Tốt/Tổng số lượng tham gia khảo sátt: {label_good_rounded}%")
        print("============================================================")
        print("************************************************************")
        print("============================================================")
    print("Proceess Complete....")
        
def copy_and_rename_file(src_path, dest_dir):
    print(f"Create File Survey...")
    # Lấy thời gian hiện tại và định dạng tên tệp
    current_time = datetime.now().strftime("%d%m%Y_%H%M")
    # Định dạng tên tệp mới (loại bỏ ký tự đặc biệt)
    new_file_name = ''.join(e for e in current_time if e.isalnum())    
    # Đường dẫn đầy đủ của tệp đích
    dest_path = os.path.join(dest_dir, f"{new_file_name}_Survey.docx")
    # Sao chép tệp
    shutil.copy(src_path, dest_path)
    print(f"File created and moved to {dest_path}")
    return dest_path

def copy_file(src_path, dest_dir):
    print(f"Checking file...")
    shutil.copy(src_path, dest_dir)
    print(f"File moved to {dest_dir}")
    return dest_dir

def move_file(src_path, dest_dir):
    # Kiểm tra xem tệp nguồn có tồn tại hay không
    if not os.path.isfile(src_path):
        print(f"Source file '{src_path}' does not exist.")
        return None

    # Kiểm tra xem thư mục đích có tồn tại hay không, nếu không thì tạo nó
    if not os.path.isdir(dest_dir):
        print(f"Destination directory '{dest_dir}' does not exist. Creating it.")
        os.makedirs(dest_dir)

    # Lấy tên tệp từ đường dẫn nguồn
    file_name = os.path.basename(src_path)
    
    # Tạo đường dẫn đích đầy đủ
    dest_path = os.path.join(dest_dir, file_name)
    
    try:
        # Di chuyển tệp
        shutil.move(src_path, dest_path)
        print(f"File moved to {dest_path}")
        return dest_path
    except Exception as e:
        print(f"Error moving file: {e}")
        return None

def write_to_word1(template_path, output_directory, data):
    # Mở file Word template
    doc = Document(template_path)
    
    # Duyệt qua tất cả các đoạn văn trong tài liệu và thay thế placeholders
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(f'{{{key}}}', value)

    # Duyệt qua tất cả các bảng trong tài liệu và thay thế placeholders
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(f'{{{key}}}', value)
    
    # Định dạng lại tên file với ngày tháng năm giờ phút
    current_time = datetime.now().strftime("%d%m%Y_%H%M")
    output_filename = f"Survey_{current_time}.docx"
    output_path = os.path.join(output_directory, output_filename)
    
    # Lưu file Word mới với dữ liệu đã thay thế
    doc.save(output_path)
    
    print(f"File đã được ghi dữ liệu và lưu tại: {output_path}")
    return output_path

def data_list_survey():
    list_survey = [
        {
            "Vui lòng đánh giá Năng lực chuyên môn?": {
                "CM_VOTE0": { "CM_RATE0": "", "CM_VALUE0":  ""},
                "CM_VOTE1": { "CM_RATE1": "", "CM_VALUE1": ""},
                "CM_VOTE2": { "CM_RATE2": "", "CM_VALUE2": ""},
                "CM_VOTE3": { "CM_RATE3": "", "CM_VALUE3": ""},
                "CM_VOTE4": { "CM_RATE4": "", "CM_VALUE4": ""},
                "CM_VOTE5": { "CM_RATE5": "", "CM_VALUE5": ""},
            }
        }]
    
    data = {
            "CM_RATE0": "7/16",
            "CM_VALUE0": "80%",
            "CM_RATE1": "3/16",
            "CM_VALUE1": "15%",
            "CM_RATE2": "3/16",
            "CM_VALUE2": "15%",
            "CM_RATE3": "3/16",
            "CM_VALUE3": "15%",
            "CM_RATE4": "25%",
            "CM_VALUE4": "35%",
        }
     
    for survey_item in list_survey:
        for key, value in survey_item.items():
            list_key = list(value)
            list_value = value
            # for i in {1,2,3,4,5}:
                # a[i] in list_sv.items():

            for vote_key, vote_value in list_value.items():
                vote_value[1] = "7/16"
                vote_value[2] = "80%"
            print(key)
            print(a[0])
            print(list_survey)

def data_list():
    list_survey = [
        {
            "Vui lòng đánh giá Năng lực chuyên môn?": {
                "CM_VOTE0": { "CM_RATE0": "", "CM_VALUE0": ""},
                "CM_VOTE1": { "CM_RATE1": "", "CM_VALUE1": ""},
                "CM_VOTE2": { "CM_RATE2": "", "CM_VALUE2": ""},
                "CM_VOTE3": { "CM_RATE3": "", "CM_VALUE3": ""},
                "CM_VOTE4": { "CM_RATE4": "", "CM_VALUE4": ""},
                "CM_VOTE5": { "CM_RATE5": "", "CM_VALUE5": ""},
            }
        }
    ]
    
    data = {
        "CM_RATE0": "",
        "CM_VALUE0": "",
        "CM_RATE1": "",
        "CM_VALUE1": "",
        "CM_RATE2": "",
        "CM_VALUE2": "",
        "CM_RATE3": "",
        "CM_VALUE3": "",
        "CM_RATE4": "",
        "CM_VALUE4": "",
    }

    data_survey = [
           {'0': {'0': '1/16', '1': 6.25}}, 
           {'1': {'0': '7/16', '1': 43.75}},
           {'2': {'0': '6/16', '1': 37.5}}
        ]
    i_value = 0
    for survey_item in list_survey:
            for key, value in survey_item.items():
                list_key = list(value.keys())
                for list_key[i_value] in value:
                    print(list_key[i_value])
                    print(i_value)
                    vote_rate = list(value[list_key[i_value]])
                    for i in range(len(vote_rate)):
                        print(i_value)
                        if vote_rate[i] in data:
                            for key_survey,value_survey in data_survey[i_value].items():
                                print(i_value)
                                for inner_key, inner_value in value_survey.items():
                                    if i == int(inner_key):
                                        data[vote_rate[i]] = inner_value  
                    i_value += 1
                            # for inner_key, inner_value in value_su.items():
                            #     if i == int(inner_key):
                            #         data[vote_rate[i]] = inner_value
    
    # for survey in data_survey:
    #    for key_su, value_su in survey.items():
    #     no_val = int(key_su)
    #     print(f"Key: {key_su}, Value: {value_su}")
        

    # for survey_item in list_survey:
    #     no_value = 0
    #     for key, value in survey_item.items():
    #         list_key = list(value.keys())
    #         val = 100
    #         for list_key[no_value] in value:
    #             vote_rate = list(value[list_key[no_value]])
    #             for i in range(len(vote_rate)):
    #                 if vote_rate[i] in data:
    #                     data[vote_rate[i]] = str(val)+"%"
    #             val += 10

    #             print(data)
    
    print(data)

def read_survey_Check(driver,all_reponse = 16):
    print("Proceess Loading....")
    list_survey = [
        {
            "Vui lòng đánh giá Năng lực chuyên môn?": {
                "CM_VOTE0": { "CM_RATE0": "", "CM_VALUE0": ""},
                "CM_VOTE1": { "CM_RATE1": "", "CM_VALUE1": ""},
                "CM_VOTE2": { "CM_RATE2": "", "CM_VALUE2": ""},
                "CM_VOTE3": { "CM_RATE3": "", "CM_VALUE3": ""},
                "CM_VOTE4": { "CM_RATE4": "", "CM_VALUE4": ""},
                "CM_VOTE5": { "CM_RATE5": "", "CM_VALUE5": ""},
            }
        }
    ]

    data = {
        "CM_RATE0":  "",
        "CM_VALUE0": "",
        "CM_RATE1":  "",
        "CM_VALUE1": "",
        "CM_RATE2":  "",
        "CM_VALUE2": "",
        "CM_RATE3":  "",
        "CM_VALUE3": "",
        "CM_RATE4":  "",
        "CM_VALUE4": "",
    }

    data_list = []

    # Duyệt qua từng danh mục và in tên của từng danh mục
    time.sleep(3)
    response_element = WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.XPATH, '//div[contains(@aria-label, "responses")]'))
    )
    # Lấy giá trị số lượng responses
    # response_value = response_element.text.strip()
    response_value = response_element.text.strip().split('\n')[-1]
    empty_value = all_reponse - int(response_value)
    print(f'Số lượng tham gia khảo Sát:{response_value}')

    
    for survey_item in list_survey:
        good_value = 0
        good_rate = 0
        no_value = 0 # new
        for key,key_value in survey_item.items():
            list_key = list(key_value.keys()) #new
            # for list_key[no_value] in value:
            #     vote_rate = list(value[list_key[no_value]])
            #     for i in range(len(vote_rate)):
            #         if vote_rate[i] in data:
            #             data[vote_rate[i]] = ""
            print(key)
            print("------------------------------------------------------------")
            xpath_query = f'//div[contains(@aria-label, "{key}")]//tbody//tr'
            rows = driver.find_elements(By.XPATH, xpath_query)
        
            for row in rows:
                label = row.find_element(By.CLASS_NAME, 'chart-control-legend-label').text
                value = row.find_element(By.CLASS_NAME, 'chart-control-legend-value').text
                vote_value = int(value)/all_reponse*100
                print(f"{label}: {value} - {vote_value}")
                # if label == "Tốt" or label == "Rất Tốt" or label == "tốt" or label == "Rất tốt":
                # if label.lower() == label_t.lower() or label.lower() == label_rt.lower():
                if label.lower() in {"tốt", "rất tốt"}:
                    good_value += vote_value
                    good_rate += int(value)

                json_data = {
                    f"{no_value}": {
                        "0": str(value)+"/"+str(all_reponse),
                        "1": vote_value
                    }
                }
                data_list.append(json_data)
                no_value += 1

                # if no_value < 3:
                #     calculate_rate(list_key,key_value, no_value, data,value,vote_value )
                #     no_value += 1
        # print(data_list)
        # calculate_rate(list_key,key_value, data,data_list )

        lable_empty = empty_value/all_reponse*100
        print(f"Không Khảo Sát: {empty_value} - {lable_empty}%")
        json_data = {
                    f"{no_value}": {
                        "0": str(empty_value)+"/"+str(all_reponse),
                        "1": lable_empty
                    }
                }
        data_list.append(json_data)
        no_value += 1

        print(f"Tỷ lệ Rất tốt và Tốt/Tổng số lượng mời tham gia khảo sát: {good_value}%")
        print(f"Không Khảo Sát: {empty_value} - {lable_empty}%")
        json_data = {
                    f"{no_value}": {
                        "0": "",
                        "1": good_value
                    }
                }
        data_list.append(json_data)
        no_value += 1
        
        label_good = good_rate/int(response_value)*100
        label_good_rounded = round(label_good, 2)
        print(f"Tỷ lệ Rất tốt và Tốt/Tổng số lượng tham gia khảo sátt: {label_good_rounded}%")
        json_data = {
                    f"{no_value}": {
                        "0": "",
                        "1": label_good_rounded
                    }
                }
        data_list.append(json_data)
        print(data_list)
        calculate_rate(list_survey, list_key,key_value, data, data_list)
        no_value = 0
        print(data)

        print("============================================================")
        print("************************************************************")
        print("============================================================")
    print("Proceess Complete....")
    print(data)

def read_survey(driver,all_reponse):
    print("Proceess Loading....")
    list_survey = [
        {
            "Vui lòng đánh giá Năng lực chuyên môn?": {
                "CM_VOTE0": { "CM_RATE0": "", "CM_VALUE0": ""},
                "CM_VOTE1": { "CM_RATE1": "", "CM_VALUE1": ""},
                "CM_VOTE2": { "CM_RATE2": "", "CM_VALUE2": ""},
                "CM_VOTE3": { "CM_RATE3": "", "CM_VALUE3": ""},
                "CM_VOTE4": { "CM_RATE4": "", "CM_VALUE4": ""},
                "CM_VOTE5": { "CM_RATE5": "", "CM_VALUE5": ""},
            }
        },
        {
            "Vui lòng đánh giá Năng lực lãnh đạo?": {
                "LD_VOTE0": { "LD_RATE0": "", "LD_VALUE0": ""},
                "LD_VOTE1": { "LD_RATE1": "", "LD_VALUE1": ""},
                "LD_VOTE2": { "LD_RATE2": "", "LD_VALUE2": ""},
                "LD_VOTE3": { "LD_RATE3": "", "LD_VALUE3": ""},
                "LD_VOTE4": { "LD_RATE4": "", "LD_VALUE4": ""},
                "LD_VOTE5": { "LD_RATE5": "", "LD_VALUE5": ""},
            }
        },
        {
            "Vui lòng đánh giá Năng lực làm việc đồng đội?": {
                "DD_VOTE0": { "DD_RATE0": "", "DD_VALUE0": ""},
                "DD_VOTE1": { "DD_RATE1": "", "DD_VALUE1": ""},
                "DD_VOTE2": { "DD_RATE2": "", "DD_VALUE2": ""},
                "DD_VOTE3": { "DD_RATE3": "", "DD_VALUE3": ""},
                "DD_VOTE4": { "DD_RATE4": "", "DD_VALUE4": ""},
                "DD_VOTE5": { "DD_RATE5": "", "DD_VALUE5": ""},
            }
        },
        {
            "Đạo đức (Liêm chính; Chính trực; Trung thành; Tận tuỵ; Lòng biết ơn)": {
                "DU_VOTE0": { "DU_RATE0": "", "DU_VALUE0": ""},
                "DU_VOTE1": { "DU_RATE1": "", "DU_VALUE1": ""},
                "DU_VOTE2": { "DU_RATE2": "", "DU_VALUE2": ""},
                "DU_VOTE3": { "DU_RATE3": "", "DU_VALUE3": ""},
                "DU_VOTE4": { "DU_RATE4": "", "DU_VALUE4": ""},
                "DU_VOTE5": { "DU_RATE5": "", "DU_VALUE5": ""},
            }
        },
        {
            "Tuân thủ kỷ luật (Tuân thủ Nội quy lao động; Tuân thủ Điều lệ; Quy chế quản lý nội bộ; Quy trình làm việc)": {
                "KL_VOTE0": { "KL_RATE0": "", "KL_VALUE0": ""},
                "KL_VOTE1": { "KL_RATE1": "", "KL_VALUE1": ""},
                "KL_VOTE2": { "KL_RATE2": "", "KL_VALUE2": ""},
                "KL_VOTE3": { "KL_RATE3": "", "KL_VALUE3": ""},
                "KL_VOTE4": { "KL_RATE4": "", "KL_VALUE4": ""},
                "KL_VOTE5": { "KL_RATE5": "", "KL_VALUE5": ""},
            }
        },
        {
            "Nhiệt huyết": {
                "NH_VOTE0": { "NH_RATE0": "", "NH_VALUE0": ""},
                "NH_VOTE1": { "NH_RATE1": "", "NH_VALUE1": ""},
                "NH_VOTE2": { "NH_RATE2": "", "NH_VALUE2": ""},
                "NH_VOTE3": { "NH_RATE3": "", "NH_VALUE3": ""},
                "NH_VOTE4": { "NH_RATE4": "", "NH_VALUE4": ""},
                "NH_VOTE5": { "NH_RATE5": "", "NH_VALUE5": ""},
            }
        },
        {
            "Lắng nghe, tôn trọng và đối xử công bằng": {
                "LN_VOTE0": { "LN_RATE0": "", "LN_VALUE0": ""},
                "LN_VOTE1": { "LN_RATE1": "", "LN_VALUE1": ""},
                "LN_VOTE2": { "LN_RATE2": "", "LN_VALUE2": ""},
                "LN_VOTE3": { "LN_RATE3": "", "LN_VALUE3": ""},
                "LN_VOTE4": { "LN_RATE4": "", "LN_VALUE4": ""},
                "LN_VOTE5": { "LN_RATE5": "", "LN_VALUE5": ""},
            }
        }
    ]

    data = {
        #Vui lòng đánh giá Năng lực chuyên môn?
        "CM_RATE0":  "", #chưa tốt
        "CM_VALUE0": "", #chưa tốt
        "CM_RATE1":  "", #tốt
        "CM_VALUE1": "", #tốt
        "CM_RATE2":  "", #rất tốt
        "CM_VALUE2": "", #rất tốt
        "CM_RATE3":  "", #không khảo sát
        "CM_VALUE3": "", #không khảo sát
        "CM_VALUE4": "", #Tỷ lệ Rất tốt và Tốt/Tổng số lượng mời tham gia khảo sát
        "CM_VALUE5": "", #Tỷ lệ Rất tốt và Tốt/Tổng số lượng tham gia khảo sát

        #Vui lòng đánh giá Năng lực lãnh đạo
        "LD_RATE0":  "", #chưa tốt
        "LD_VALUE0": "", #chưa tốt
        "LD_RATE1":  "", #tốt
        "LD_VALUE1": "", #tốt
        "LD_RATE2":  "", #rất tốt
        "LD_VALUE2": "", #rất tốt
        "LD_RATE3":  "", #không khảo sát
        "LD_VALUE3": "", #không khảo sát
        "LD_VALUE4": "", #Tỷ lệ Rất tốt và Tốt/Tổng số lượng mời tham gia khảo sát
        "LD_VALUE5": "", #Tỷ lệ Rất tốt và Tốt/Tổng số lượng tham gia khảo sát

        #Vui lòng đánh giá Năng lực làm việc đồng đội?
        "DD_RATE0":  "", #chưa tốt
        "DD_VALUE0": "", #chưa tốt
        "DD_RATE1":  "", #tốt
        "DD_VALUE1": "", #tốt
        "DD_RATE2":  "", #rất tốt
        "DD_VALUE2": "", #rất tốt
        "DD_RATE3":  "", #không khảo sát
        "DD_VALUE3": "", #không khảo sát
        "DD_VALUE4": "", #Tỷ lệ Rất tốt và Tốt/Tổng số lượng mời tham gia khảo sát
        "DD_VALUE5": "", #Tỷ lệ Rất tốt và Tốt/Tổng số lượng tham gia khảo sát

        #Đạo đức (Liêm chính; Chính trực; Trung thành; Tận tuỵ; Lòng biết ơn)?
        "DU_RATE0":  "", #chưa tốt
        "DU_VALUE0": "", #chưa tốt
        "DU_RATE1":  "", #tốt
        "DU_VALUE1": "", #tốt
        "DU_RATE2":  "", #rất tốt
        "DU_VALUE2": "", #rất tốt
        "DU_RATE3":  "", #không khảo sát
        "DU_VALUE3": "", #không khảo sát
        "DU_VALUE4": "", #Tỷ lệ Rất tốt và Tốt/Tổng số lượng mời tham gia khảo sát
        "DU_VALUE5": "", #Tỷ lệ Rất tốt và Tốt/Tổng số lượng tham gia khảo sát

        #Tuân thủ kỷ luật (Tuân thủ Nội quy lao động; Tuân thủ Điều lệ; Quy chế quản lý nội bộ; Quy trình làm việc)
        "KL_RATE0":  "", #chưa tốt
        "KL_VALUE0": "", #chưa tốt
        "KL_RATE1":  "", #tốt
        "KL_VALUE1": "", #tốt
        "KL_RATE2":  "", #rất tốt
        "KL_VALUE2": "", #rất tốt
        "KL_RATE3":  "", #không khảo sát
        "KL_VALUE3": "", #không khảo sát
        "KL_VALUE4": "", #Tỷ lệ Rất tốt và Tốt/Tổng số lượng mời tham gia khảo sát
        "KL_VALUE5": "", #Tỷ lệ Rất tốt và Tốt/Tổng số lượng tham gia khảo sát

        #Nhiệt huyết
        "NH_RATE0":  "", #chưa tốt
        "NH_VALUE0": "", #chưa tốt
        "NH_RATE1":  "", #tốt
        "NH_VALUE1": "", #tốt
        "NH_RATE2":  "", #rất tốt
        "NH_VALUE2": "", #rất tốt
        "NH_RATE3":  "", #không khảo sát
        "NH_VALUE3": "", #không khảo sát
        "NH_VALUE4": "", #Tỷ lệ Rất tốt và Tốt/Tổng số lượng mời tham gia khảo sát
        "NH_VALUE5": "", #Tỷ lệ Rất tốt và Tốt/Tổng số lượng tham gia khảo sát

        #Lắng nghe, tôn trọng và đối xử công bằng
        "LN_RATE0":  "", #chưa tốt
        "LN_VALUE0": "", #chưa tốt
        "LN_RATE1":  "", #tốt
        "LN_VALUE1": "", #tốt
        "LN_RATE2":  "", #rất tốt
        "LN_VALUE2": "", #rất tốt
        "LN_RATE3":  "", #không khảo sát
        "LN_VALUE3": "", #không khảo sát
        "LN_VALUE4": "", #Tỷ lệ Rất tốt và Tốt/Tổng số lượng mời tham gia khảo sát
        "LN_VALUE5": "", #Tỷ lệ Rất tốt và Tốt/Tổng số lượng tham gia khảo sát
    }

    data_list = []
    # Duyệt qua từng danh mục và in tên của từng danh mục
    time.sleep(3)
    response_element = WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.XPATH, '//div[contains(@aria-label, "responses")]'))
    )
    # Lấy giá trị số lượng responses
    # response_value = response_element.text.strip()
    response_value = response_element.text.strip().split('\n')[-1]
    empty_value = all_reponse - int(response_value)
    print(f'Số lượng tham gia khảo Sát:{response_value}')

    list_survey_no = 0
    for survey_item in list_survey:
        good_value = 0
        good_rate = 0
        no_value = 0 # new
        for key,key_value in survey_item.items():
            list_key = list(key_value.keys()) #new
            print(key)
            print("------------------------------------------------------------")
            xpath_query = f'//div[contains(@aria-label, "{key}")]//tbody//tr'
            rows = driver.find_elements(By.XPATH, xpath_query)
        
            for row in rows:
                label = row.find_element(By.CLASS_NAME, 'chart-control-legend-label').text
                value = row.find_element(By.CLASS_NAME, 'chart-control-legend-value').text
                vote_value = int(value)/all_reponse*100
                print(f"{label}: {value} - {vote_value}")
                # if label == "Tốt" or label == "Rất Tốt" or label == "tốt" or label == "Rất tốt":
                # if label.lower() == label_t.lower() or label.lower() == label_rt.lower():
                if label.lower() in {"tốt", "rất tốt"}:
                    good_value += vote_value
                    good_rate += int(value)

                json_data = {
                    f"{no_value}": {
                        "0": str(value)+"/"+str(all_reponse),
                        # "1": str(vote_value)+"%"
                        "1": str(round(vote_value, 2))+"%"
                    }
                }
                data_list.append(json_data)
                no_value += 1

        lable_empty = empty_value/all_reponse*100
        print(f"Không Khảo Sát: {empty_value} - {lable_empty}%")
        json_data = {
                    f"{no_value}": {
                        "0": str(empty_value)+"/"+str(all_reponse),
                        # "1": str(lable_empty)+"%"
                        "1": str(round(lable_empty, 2))+"%"
                    }
                }
        data_list.append(json_data)
        no_value += 1

        print(f"Tỷ lệ Rất tốt và Tốt/Tổng số lượng mời tham gia khảo sát: {good_value}%")
        # print(f"Không Khảo Sát: {empty_value} - {lable_empty}%")
        json_data = {
                    f"{no_value}": {
                        "0": "",
                        # "1": str(good_value)+"%"
                        "1": str(round(good_value, 2))+"%"
                    }
                }
        data_list.append(json_data)
        no_value += 1
        
        label_good = good_rate/int(response_value)*100
        label_good_rounded = round(label_good, 2)
        print(f"Tỷ lệ Rất tốt và Tốt/Tổng số lượng tham gia khảo sátt: {label_good_rounded}%")
        json_data = {
                    f"{no_value}": {
                        "0": "",
                        "1": str(label_good_rounded)+"%"
                    }
                }
        data_list.append(json_data)
        # print(data_list)
        calculate_rate(list_survey,list_survey_no, list_key,key_value, data, data_list)
        no_value = 0
        # print(data)
        list_survey_no += 1

        print("============================================================")
        print("************************************************************")
        print("============================================================")
    print("Proceess Complete....")
    # print(data)
    return data

def calculate_rate(list_survey,list_survey_no, list_key,key_value, data, data_list):
    # list_key = list(value.keys()) #new
    try:
        i_value = 0
        survey_item = list_survey[list_survey_no]
        # for survey_key, survey_item in list_survey[list_survey_no]:
        for key, value in survey_item.items():
            list_key = list(value.keys())
            for list_key[i_value] in value:
                # print(list_key[i_value])
                # print(i_value)
                vote_rate = list(value[list_key[i_value]])
                for i in range(len(vote_rate)):
                    # print(i_value)
                    if vote_rate[i] in data:
                        for key_survey,value_survey in data_list[i_value].items():
                            # print(i_value)
                            for inner_key, inner_value in value_survey.items():
                                if i == int(inner_key):
                                    data[vote_rate[i]] = inner_value  
                i_value += 1
        # print(json.dumps(data, indent=4, ensure_ascii=False))
        data_list.clear()
    except Exception as e:
        print(e)

def convert_to_pdf(docx_path):
    pdf_path = docx_path.replace(".docx", ".pdf")
    convert(docx_path, pdf_path)
    return pdf_path

def open_file(canvas, interior):
    file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")])
    if file_path:
        if file_path.endswith(".docx"):
            pdf_path = convert_to_pdf(file_path)
        else:
            pdf_path = file_path
        show_pdf(pdf_path, canvas, interior)

def show_pdf(file_path, canvas, root):
    try:
        doc = fitz.open(file_path)
        num_pages = doc.page_count

        # Xóa trước khi hiển thị trang mới
        canvas.delete("all")

        # Lấy kích thước của trang đầu tiên để điều chỉnh kích thước cửa sổ
        first_page = doc.load_page(0)
        first_pix = first_page.get_pixmap(alpha=False)
        width, height = first_pix.width, first_pix.height

        # Điều chỉnh kích thước cửa sổ
        root.geometry(f"{width}x{height + 100}")  # +100 để có không gian cho scrollbar

        # Hiển thị từng trang dưới dạng ảnh và cuộn xuống
        y_position = 0
        for i in range(num_pages):
            page = doc.load_page(i)
            pix = page.get_pixmap(alpha=False)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            photo = ImageTk.PhotoImage(img)

            # Thêm hình ảnh vào canvas
            canvas.create_image(0, y_position, anchor="nw", image=photo)
            y_position += pix.height

            # Giữ tham chiếu để tránh mất ảnh
            if not hasattr(canvas, 'images'):
                canvas.images = []
            canvas.images.append(photo)

        # Cập nhật kích thước nội dung của canvas
        canvas.config(scrollregion=canvas.bbox("all"))

    except Exception as e:
        messagebox.showerror("Error", str(e))

def show_pdf_in_tkinter(filepath=None):
    # Kiểm tra phần mềm PDF viewer trên máy
    # if os.system("which acroread") == 0:  # Sử dụng Adobe Reader nếu có
    #     os.system(f"acroread {filepath}")
    #     return
    # elif os.system("which evince") == 0:  # Sử dụng Evince nếu có
    #     os.system(f"evince {filepath}")
    #     return

    # Tạo cửa sổ GUI
    root = tk.Tk()
    root.title("PDF Viewer")

    # Tạo khung chứa PDF
    frame = tk.Frame(root)
    frame.pack(fill='both', expand=True)

    # Tạo canvas để cuộn
    canvas = tk.Canvas(frame)
    canvas.pack(side="left", fill="both", expand=True)

    # Thêm thanh cuộn dọc
    scrollbar = tk.Scrollbar(frame, orient="vertical", command=canvas.yview)
    scrollbar.pack(side="right", fill="y")

    # Cấu hình canvas với scrollbar
    canvas.config(yscrollcommand=scrollbar.set)

    # Tạo menu
    menu = tk.Menu(root)
    root.config(menu=menu)

    # Thêm tùy chọn "File" vào menu
    file_menu = tk.Menu(menu, tearoff=0)
    menu.add_cascade(label="File", menu=file_menu)
    file_menu.add_command(label="Open", command=lambda: open_file(canvas, root))
    file_menu.add_separator()
    file_menu.add_command(label="Exit", command=root.quit)

    # Mở tệp PDF nếu filepath được cung cấp
    if filepath:
        show_pdf(filepath, canvas, root)

    # Chạy ứng dụng GUI
    root.mainloop()

def replace_text_with_format(paragraph, key, value):
    # Tạo một danh sách mới để chứa các đoạn văn bản đã thay thế
    new_paragraph_runs = []

    for run in paragraph.runs:
        if key in run.text:
            # Chia nhỏ các đoạn văn bản bằng khóa
            parts = run.text.split(key)
            for i, part in enumerate(parts):
                # Tạo một đoạn văn bản mới với phần trước và sau khóa
                if part:
                    new_run = run._element
                    new_run.text = part
                    new_paragraph_runs.append(new_run)
                if i < len(parts) - 1:
                    # Chèn giá trị thay thế với định dạng tương tự
                    replacement_run = run._element
                    replacement_run.text = value
                    new_paragraph_runs.append(replacement_run)
        else:
            # Giữ nguyên đoạn văn bản nếu không chứa khóa
            new_paragraph_runs.append(run._element)

    # Xóa tất cả các đoạn văn bản hiện có trong đoạn văn
    while paragraph.runs:
        paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)

    # Thêm các đoạn văn bản mới vào đoạn văn
    for run_element in new_paragraph_runs:
        paragraph._element.append(run_element)

def write_to_word(template_path, output_directory, data):
    # Mở file Word template
    doc = Document(template_path)

    # Duyệt qua tất cả các đoạn văn trong tài liệu và thay thế placeholders
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                if any(run.bold or run.italic or run.underline or run.font.name or (run.font.size is not None) for run in paragraph.runs):
                    replace_text_with_format(paragraph, f'{{{key}}}', value)
                else:
                    paragraph.text = paragraph.text.replace(f'{{{key}}}', value)

    # Duyệt qua tất cả các bảng trong tài liệu và thay thế placeholders
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        if key in paragraph.text:
                            if any(run.bold or run.italic or run.underline or run.font.name or (run.font.size is not None) for run in paragraph.runs):
                                replace_text_with_format(paragraph, f'{{{key}}}', value)
                            else:
                                paragraph.text = paragraph.text.replace(f'{{{key}}}', value)
    
    # Định dạng lại tên file với ngày tháng năm giờ phút
    current_time = datetime.now().strftime("%d%m%Y_%H%M")
    output_filename = f"Survey_{current_time}.docx"
    output_path = os.path.join(output_directory, output_filename)
    
    # Lưu file Word mới với dữ liệu đã thay thế
    doc.save(output_path)
    
    print(f"File đã được ghi dữ liệu và lưu tại: {output_path}")
    return output_path

if __name__ == "__main__":
    # Yêu cầu người dùng nhập đường dẫn tới tệp Excel
    # file_path = input("Vui lòng nhập tổng số người tham gia khảo sát: ")
    # file_path = input("Vui lòng nhập đường dẫn tới tệp Link Survey: ")
    # file_path = "https://forms.office.com/Pages/DesignPageV2.aspx?prevorigin=shell&origin=NeoPortalPage&subpage=design&id=gT89xHr1zEiLB3SzmTXYdopPyoq-8NZGhRvMGc1_G9pUMlVDNFc3SEZQVVhPRkdHREJYSFI0S0Q5Ti4u&analysis=true"

    # Đọc dữ liệu từ tệp Excel
    while True:
        # Yêu cầu người dùng nhập tổng số người tham gia khảo sát
        total_participants = input("Vui lòng nhập tổng số người tham gia khảo sát: ")
        if total_participants:
            break
        print("Tổng số người tham gia khảo sát không được để trống. Vui lòng nhập lại.")

    while True:
        # Yêu cầu người dùng nhập đường dẫn tới tệp Link Survey
        file_path = input("Vui lòng nhập đường dẫn tới tệp Link Survey: ")
        if file_path:
            break
        print("Đường dẫn tới tệp Link Survey không được để trống. Vui lòng nhập lại.")

    try:
        driver = setup_firefox_driver()
        login_to_website(driver, file_path)
        data = read_survey(driver, int(total_participants))
        # print(json.dumps(data, indent=4, ensure_ascii=False))
        # Đường dẫn tệp nguồn
        src_file_template = ".\Template\Template_Survey.docx"
        dest_directory = ".\Survey"
        src_output_survey = write_to_word(src_file_template, dest_directory, data)
        pdf_file = convert_to_pdf(src_output_survey)
        src_pdf = ".\PDF"
        dest_path = move_file(pdf_file, src_pdf)
        show_pdf_in_tkinter(dest_path)
    except ValueError:
        print("Error")

    input("Nhấn phím bất kỳ để kết thúc...")

